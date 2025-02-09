import os
import re
import json
import cv2
import numpy as np
import pytesseract
import spacy
import matplotlib.pyplot as plt
import seaborn as sns
from pdf2image import convert_from_path
from docx import Document
from PIL import Image
import pdfplumber
from sklearn.cluster import DBSCAN

# Load spaCy large model
nlp = spacy.load("en_core_web_lg")

# Pre-compile a regex for token extraction.
WORD_REGEX = re.compile(r'\w+')

###############################################################################
# 1. SECTION SYNONYMS SETUP (Expanded with Singular Forms)
###############################################################################
RAW_SECTION_SYNONYMS = {
    "WORK EXPERIENCE": ["work experience", "experience", "employment history", "professional experience", "career history"],
    "EDUCATION": ["education", "academic background", "academic qualification", "study", "qualifications", "educational history"],
    "SKILLS": ["skills", "technologies", "technical skills", "core competencies", "competencies", "abilities"],
    "PROJECTS": ["projects", "portfolio", "assignments", "research", "case studies"],
    "SUMMARY": ["summary", "overview", "profile summary", "about me", "personal profile"],
    "ACHIEVEMENTS": ["achievements", "accomplishments", "awards", "milestones", "achievement"],
    "CERTIFICATIONS": ["certifications", "certificates", "training", "certification"],
    "LANGUAGES": ["languages", "spoken languages", "linguistic abilities"],
    "POSITIONS OF RESPONSIBILITY": ["positions of responsibility", "leadership roles", "extracurricular activities", "volunteer work"],
    "REFERENCES": ["references", "referees", "recommendations"],
    "INTERNSHIPS": ["internships", "industrial training", "work placements"],
    "PUBLICATIONS": ["publications", "research papers", "articles"],
    "HOBBIES": ["hobbies", "interests", "personal interests"],
}

def singular(word):
    """Naively converts a word to singular form by removing a trailing 's'."""
    return word[:-1] if word.endswith('s') else word

SECTION_SYNONYMS = {}
for canonical, syn_list in RAW_SECTION_SYNONYMS.items():
    expanded = set()
    for phrase in syn_list:
        norm_phrase = phrase.lower().strip()
        words = norm_phrase.split()
        expanded_words = []
        for w in words:
            expanded_words.append(w)
            expanded_words.append(singular(w))
        expanded.add(norm_phrase)
        expanded.add(" ".join(expanded_words))
    SECTION_SYNONYMS[canonical] = list(expanded)

###############################################################################
# 2. TEXT NORMALIZATION & TOKEN MAPPING
###############################################################################
def normalize_text(text):
    """
    Lowercases the text, strips whitespace, and removes punctuation.
    For example, "Skills & Awards" becomes "skills awards".
    """
    text = text.lower().strip()
    text = re.sub(r'[^\w\s]', '', text)
    return text

def token_to_canonical(token):
    """
    Given a token, returns the canonical header category if this token exactly appears
    in any of the section_synonyms.
    """
    for canonical, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            if token == normalize_text(syn):
                return canonical
    return None

###############################################################################
# 3. SPELLING CORRECTION (for OCR errors)
###############################################################################
def correct_spelling(text):
    """Corrects common OCR errors using a simple mapping."""
    corrections = {
        "skiel": "skill",
        "skils": "skill",
    }
    words = text.split()
    corrected_words = [corrections.get(word.lower(), word) for word in words]
    return " ".join(corrected_words)

###############################################################################
# 4. MERGE OVERLAPPING BOXES
###############################################################################
def merge_overlapping_boxes(sections, tol=2):
    """Merges detected text region boxes that overlap or are very close."""
    if not sections:
        return []
    sections = sorted(sections, key=lambda s: s.get("y", 0))
    merged = []
    used = [False] * len(sections)
    for i in range(len(sections)):
        if used[i]:
            continue
        box1 = sections[i]
        x1, y1, w1, h1 = box1["x"], box1["y"], box1["w"], box1["h"]
        x2, y2 = x1 + w1, y1 + h1
        merged_box = box1.copy()
        used[i] = True
        for j in range(i+1, len(sections)):
            if used[j]:
                continue
            box2 = sections[j]
            if (box2["x"] >= x1 - tol and box2["x"] <= x2 + tol and
                box2["y"] >= y1 - tol and box2["y"] <= y2 + tol):
                nx1 = min(x1, box2["x"])
                ny1 = min(y1, box2["y"])
                nx2 = max(x2, box2["x"] + box2["w"])
                ny2 = max(y2, box2["y"] + box2["h"])
                x1, y1, x2, y2 = nx1, ny1, nx2, ny2
                merged_box["x"] = x1
                merged_box["y"] = y1
                merged_box["w"] = x2 - x1
                merged_box["h"] = ny2 - y1
                merged_box["text"] = (merged_box.get("text", "") + "\n" + box2.get("text", "")).strip()
                used[j] = True
        merged.append(merged_box)
    return merged

###############################################################################
# 4A. CLUSTERING OF BOUNDING BOXES USING DBSCAN
###############################################################################
def cluster_bounding_boxes(sections, eps=20, min_samples=1):
    """
    Uses DBSCAN to cluster the centers of the bounding boxes.
    Boxes in the same cluster are merged into one.
    Returns a new dictionary of merged sections.
    """
    centers = []
    keys = []
    for key, sec in sections.items():
        x, y, w, h = sec.get("x", 0), sec.get("y", 0), sec.get("w", 0), sec.get("h", 0)
        centers.append([x + w/2, y + h/2])
        keys.append(key)
    
    if not centers:
        return sections

    dbscan = DBSCAN(eps=eps, min_samples=min_samples)
    labels = dbscan.fit_predict(np.array(centers))
    
    clusters = {}
    for key, label in zip(keys, labels):
        clusters.setdefault(label, []).append(key)
    
    merged_sections = {}
    new_index = 1
    for label, key_list in clusters.items():
        if len(key_list) == 1:
            merged_sections[new_index] = sections[key_list[0]]
            new_index += 1
        else:
            min_x = min(sections[k]["x"] for k in key_list)
            min_y = min(sections[k]["y"] for k in key_list)
            max_x = max(sections[k]["x"] + sections[k]["w"] for k in key_list)
            max_y = max(sections[k]["y"] + sections[k]["h"] for k in key_list)
            merged_text = "\n".join(sections[k].get("text", "") for k in key_list)
            is_header = any(sections[k].get("is_header") for k in key_list)
            canonical_header = None
            for k in key_list:
                if sections[k].get("is_header"):
                    canonical_header = sections[k].get("canonical_header")
                    break
            merged_sections[new_index] = {
                "x": min_x,
                "y": min_y,
                "w": max_x - min_x,
                "h": max_y - min_y,
                "text": merged_text,
                "is_header": is_header,
                "canonical_header": canonical_header
            }
            new_index += 1
    return merged_sections

###############################################################################
# 5. HEADER DETECTION (Rule-Based)
###############################################################################
def ends_with_newline_rule(line):
    """Returns True if the line starts or ends with a newline, colon, or dash."""
    return line.startswith("\n") or line.endswith("\n") or line.endswith(":") or line.endswith("-")

def exact_match_header(line):
    """Returns the canonical header if the entire normalized text exactly matches a section synonym."""
    line_norm = normalize_text(line or "")
    for canonical, syns in SECTION_SYNONYMS.items():
        for syn in syns:
            if line_norm == normalize_text(syn):
                return canonical
    return None

def universal_header_check(line):
    """
    For candidate header lines (with 1-4 words):
      - For 1 word: token must map exactly.
      - For 2 words: both tokens must map exactly.
      - For 3 or 4 words: at least 2 tokens must map or a contiguous 2-word phrase must match.
    """
    tokens = WORD_REGEX.findall(normalize_text(line or ""))
    n = len(tokens)
    if n == 0:
        return False
    if n == 1:
        return token_to_canonical(tokens[0]) is not None
    elif n == 2:
        if normalize_text(line) in [normalize_text(s) for s in sum(SECTION_SYNONYMS.values(), [])]:
            return True
        return all(token_to_canonical(t) is not None for t in tokens)
    elif n in (3, 4):
        contiguous_match = False
        for i in range(n - 1):
            phrase = " ".join(tokens[i:i+2])
            for syns in SECTION_SYNONYMS.values():
                if phrase in [normalize_text(s) for s in syns]:
                    contiguous_match = True
                    break
            if contiguous_match:
                break
        token_matches = sum(1 for t in tokens if token_to_canonical(t) is not None)
        if n == 3:
            return contiguous_match or token_matches >= 2
        if n == 4:
            return contiguous_match or token_matches >= 2
    else:
        return False

def detect_uppercase_header(line):
    """Returns True if the line is fully uppercase and meets the universal header check."""
    return (line or "").isupper() and universal_header_check(line)

def rule_based_header_candidate(line):
    """
    Applies header detection rules:
      - Only lines with 1-4 words are considered.
      - Checks for an exact match, universal header match, uppercase, or punctuation cues.
      - Splits by "and" and checks each part.
    """
    line = line or ""
    tokens = WORD_REGEX.findall(line)
    if len(tokens) < 1 or len(tokens) > 4:
        return False, None
    canonical = exact_match_header(line)
    if canonical:
        return True, canonical
    if universal_header_check(line):
        return True, "Detected Header"
    if detect_uppercase_header(line):
        return True, "Detected Header"
    if ends_with_newline_rule(line) and line.strip():
        return True, "Detected Header"
    if "and" in line.lower():
        parts = [p.strip() for p in line.lower().split("and")]
        valid_parts = [p for p in parts if universal_header_check(p) or exact_match_header(p)]
        if valid_parts:
            return True, "Detected Header"
    return False, None

def is_header_candidate(line):
    """Wrapper for rule_based_header_candidate for lines with 1-4 words."""
    line = line or ""
    tokens = WORD_REGEX.findall(line)
    if len(tokens) < 1 or len(tokens) > 4:
        return False, None
    return rule_based_header_candidate(line)

def contains_header_token(line):
    """Checks if the line contains any known header token."""
    line = line or ""
    tokens = WORD_REGEX.findall(normalize_text(line))
    if not tokens:
        return False
    header_tokens = {"experience", "education", "skills", "projects", "summary",
                     "achievements", "certifications", "languages", "positions",
                     "references", "internships", "publications", "hobbies"}
    return any(token in header_tokens for token in tokens)

###############################################################################
# 6. FALLBACK & ADDITIONAL HEADER DETECTION PASSES
###############################################################################
def fallback_line_by_line_detection(sections):
    """
    For each text section, splits its text into lines and applies header detection.
    If a line has more than 4 words, only its first 4 words are checked.
    """
    if isinstance(sections, dict):
        sec_iter = sections.values()
    else:
        sec_iter = sections
    for sec in sec_iter:
        text = sec.get("text") or ""
        lines = text.split("\n")
        for line in lines:
            candidate, canonical = is_header_candidate(line)
            tokens = WORD_REGEX.findall(line)
            if not candidate and len(tokens) > 4:
                first_four = " ".join(tokens[:4])
                candidate, canonical = is_header_candidate(first_four)
            if candidate or contains_header_token(line):
                sec["is_header"] = True
                sec["canonical_header"] = canonical or "Detected Header"
                break
    return sections

def additional_header_passes(sections, num_passes=5):
    """Runs fallback_line_by_line_detection repeatedly."""
    for _ in range(num_passes):
        sections = fallback_line_by_line_detection(sections)
    return sections

def check_headers_in_tables(sections, tables):
    """
    Checks table cells (header row and left column) for headers.
    If found, adds them as new sections.
    """
    table_headers = []
    for tbl in tables:
        data = tbl.get("data", [])
        if data:
            for cell in data[0]:
                if cell and is_header_candidate(cell)[0]:
                    table_headers.append(is_header_candidate(cell)[1])
                    break
            for row in data:
                if row and row[0] and is_header_candidate(row[0])[0]:
                    table_headers.append(is_header_candidate(row[0])[1])
                    break
    for header in table_headers:
        sections[len(sections)+1] = {
            "x": 0, "y": 0, "w": 100, "h": 20,
            "text": header, "is_header": True, "canonical_header": header
        }
    return sections, table_headers

###############################################################################
# 7. COUNT HEADERS FUNCTION
###############################################################################
def count_headers(sections):
    """Counts how many sections are flagged as headers."""
    count = 0
    if isinstance(sections, dict):
        for sec in sections.values():
            if sec.get("is_header"):
                count += 1
    elif isinstance(sections, list):
        for sec in sections:
            if sec.get("is_header"):
                count += 1
    return count

###############################################################################
# 7. DOCX GROUPING FUNCTION (Order-Based)
###############################################################################
def group_docx_sections(sections, tables=None):
    """
    Groups DOCX text sections sequentially.
    The first encountered header becomes the current header; subsequent paragraphs are grouped with it.
    If no header is encountered, groups under "PROFILE".
    """
    groups = {}
    current_header_key = "PROFILE"
    groups[current_header_key] = {"header": None, "text": [], "tables": []}
    for sec in sections:
        if sec.get("is_header"):
            current_header_key = f"HEADER::{sec.get('canonical_header','')}::{sec.get('text','')}"
            groups[current_header_key] = {"header": sec, "text": [], "tables": []}
        else:
            groups[current_header_key]["text"].append(sec)
    if tables:
        for tbl in tables:
            groups[current_header_key]["tables"].append(tbl)
    return groups

###############################################################################
# 8. TABLE EXTRACTION FUNCTIONS
###############################################################################
def extract_tables_from_pdf(pdf_path):
    """Extracts tables from a PDF using pdfplumber."""
    tables_extracted = []
    with pdfplumber.open(pdf_path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            found_tables = page.find_tables()
            for table in found_tables:
                bbox = table.bbox
                table_data = table.extract()
                tables_extracted.append({
                    "page": page_number,
                    "bbox": bbox,
                    "data": table_data
                })
    return tables_extracted

def extract_tables_from_docx(docx_path):
    """Extracts tables from a DOCX file using python-docx."""
    doc = Document(docx_path)
    tables_extracted = []
    for idx, table in enumerate(doc.tables, start=1):
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables_extracted.append({
            "table_index": idx,
            "data": table_data
        })
    return tables_extracted

def extract_tables_from_image(image_path):
    """
    Uses OpenCV to detect grid-like structures that may indicate a table.
    Returns a list of dictionaries with the table's ROI and bounding box.
    """
    img = cv2.imread(image_path, cv2.IMREAD_COLOR)
    if img is None:
        return []
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    thresh = cv2.adaptiveThreshold(~gray, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 15, -2)
    horizontal = thresh.copy()
    cols = horizontal.shape[1]
    horizontal_size = cols // 30
    horizontalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (horizontal_size, 1))
    horizontal = cv2.erode(horizontal, horizontalStructure)
    horizontal = cv2.dilate(horizontal, horizontalStructure)
    vertical = thresh.copy()
    rows = vertical.shape[0]
    vertical_size = rows // 30
    verticalStructure = cv2.getStructuringElement(cv2.MORPH_RECT, (1, vertical_size))
    vertical = cv2.erode(vertical, verticalStructure)
    vertical = cv2.dilate(vertical, verticalStructure)
    mask = horizontal + vertical
    contours, _ = cv2.findContours(mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    tables = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 50 and h > 50:
            table_roi = img[y:y+h, x:x+w]
            tables.append({
                "bbox": (x, y, w, h),
                "roi": table_roi
            })
    return tables

###############################################################################
# 9. DOCX TEXT EXTRACTION (Using Updated Section_synonyms Criteria)
###############################################################################
def extract_text_from_docx(docx_path):
    """
    Extracts text from a DOCX file.
    A paragraph is flagged as a header only if its text meets the section_synonyms criteria.
    """
    print("📄 Extracting text from DOCX file...")
    doc = Document(docx_path)
    sections = []
    for para in doc.paragraphs:
        raw_text = para.text.strip()
        if not raw_text:
            continue
        candidate_txt, canonical = is_header_candidate(raw_text)
        if candidate_txt:
            sections.append({
                "text": raw_text,
                "is_header": True,
                "canonical_header": canonical,
            })
        else:
            sections.append({
                "text": raw_text,
                "is_header": False,
                "canonical_header": None,
            })
    return sections

###############################################################################
# 10. IMAGE-BASED TEXT EXTRACTION (for PDFs & Images)
###############################################################################
def generate_heatmap_and_sections(image, dilation_radius=30):
    """
    Extracts bounding boxes for text regions from a PIL Image using the specified dilation radius,
    then merges overlapping boxes.
    """
    img_cv = np.array(image)
    if len(img_cv.shape) == 2:
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_GRAY2BGR)
    elif img_cv.shape[2] == 4:
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGBA2RGB)
    gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
    edges = cv2.Canny(gray, 30, 100)
    kernel = np.ones((dilation_radius, dilation_radius), np.uint8)
    dilated = cv2.dilate(edges, kernel, iterations=1)
    heatmap = np.zeros_like(gray, dtype=np.float32)
    contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    sections = []
    for cnt in contours:
        x, y, w, h = cv2.boundingRect(cnt)
        sections.append({"x": x, "y": y, "w": w, "h": h, "is_header": False, "text": ""})
    merged_sections = merge_overlapping_boxes(sections, tol=2)
    merged_dict = {i+1: box for i, box in enumerate(merged_sections)}
    plt.figure(figsize=(10, 10))
    sns.heatmap(heatmap, cmap="hot", cbar=True)
    plt.title("Text Density Heatmap")
    plt.show()
    return merged_dict, heatmap

def ocr_sections(image, sections):
    """
    Runs OCR on each merged section from the image,
    applies spelling correction and header detection,
    and draws bounding boxes (green for headers, blue for regular text).
    """
    text_by_section = {}
    img_cv = np.array(image)
    if len(img_cv.shape) == 2:
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_GRAY2BGR)
    elif img_cv.shape[2] == 4:
        img_cv = cv2.cvtColor(img_cv, cv2.COLOR_RGBA2RGB)
    img_highlight = img_cv.copy()
    for idx, sec in sections.items():
        x, y, w, h = sec.get("x", 0), sec.get("y", 0), sec.get("w", 0), sec.get("h", 0)
        cropped = image.crop((x, y, x+w, y+h))
        text = pytesseract.image_to_string(cropped, config="--psm 6").strip()
        text = correct_spelling(text)
        if text:
            sec["text"] = text
            candidate, canonical = is_header_candidate(text)
            if candidate:
                sec["is_header"] = True
                sec["canonical_header"] = canonical
                cv2.rectangle(img_highlight, (x, y), (x+w, y+h), (0, 255, 0), 2)
            else:
                cv2.rectangle(img_highlight, (x, y), (x+w, y+h), (255, 0, 0), 2)
            text_by_section[idx] = sec
    plt.figure(figsize=(12, 8))
    plt.imshow(cv2.cvtColor(img_highlight, cv2.COLOR_BGR2RGB))
    plt.title("Highlighted Sections (Green=Header, Blue=Text)")
    plt.axis("off")
    plt.show()
    return text_by_section

###############################################################################
# 11. FILE PROCESSING FUNCTION
###############################################################################
def process_file(file_path):
    """
    Processes a PDF, DOCX, or image file for text extraction.
    Returns sections as a dict (for PDFs/images) or as a list (for DOCX).
    """
    if file_path.endswith(".docx"):
        sections = extract_text_from_docx(file_path)
        return sections
    if file_path.endswith(".pdf") or file_path.lower().endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
        if file_path.endswith(".pdf"):
            images = convert_from_path(file_path, dpi=200)
            if not images:
                print("❌ No pages found in PDF!")
                return None
            image = images[0]
        else:
            print("🖼 Processing Image File...")
            image = Image.open(file_path)
        sections, _ = generate_heatmap_and_sections(image, dilation_radius=30)
        current_dilation = 30
        found_headers = count_headers(sections)
        while current_dilation >= 5 and found_headers < 4:
            print(f"🔧 Trying dilation radius: {current_dilation}")
            sections, _ = generate_heatmap_and_sections(image, dilation_radius=current_dilation)
            sections = ocr_sections(image, sections)
            sections = fallback_line_by_line_detection(sections)
            found_headers = count_headers(sections)
            print(f"→ Found {found_headers} header(s) with dilation radius {current_dilation}")
            if found_headers >= 4:
                break
            current_dilation -= 5
        sections = additional_header_passes(sections, num_passes=5)
        # --- Cluster nearby bounding boxes to merge fragmented header parts ---
        sections = cluster_bounding_boxes(sections, eps=20, min_samples=1)
        if count_headers(sections) < 4:
            tables_dummy = []
            if file_path.endswith(".pdf"):
                tables_dummy = extract_tables_from_pdf(file_path)
            elif file_path.endswith(".docx"):
                tables_dummy = extract_tables_from_docx(file_path)
            elif file_path.lower().endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
                tables_dummy = extract_tables_from_image(file_path)
            sections, _ = check_headers_in_tables(sections, tables_dummy)
        return sections
    print("❌ Unsupported file format!")
    return None

###############################################################################
# 12. EMBEDDED LINK EXTRACTION FUNCTIONS
###############################################################################
def extract_embedded_links_from_pdf(pdf_path):
    """
    Extracts embedded hyperlinks from a PDF using PyMuPDF.
    Returns a list of unique URLs found in the link annotations.
    """
    import fitz
    doc = fitz.open(pdf_path)
    links = []
    for page in doc:
        for lnk in page.get_links():
            if "uri" in lnk:
                links.append(lnk["uri"])
    return list(set(links))

def extract_embedded_links_from_docx(docx_path):
    """
    Extracts embedded hyperlinks from a DOCX file.
    This function retrieves the underlying XML and uses regex to find href attributes.
    """
    doc = Document(docx_path)
    xml_str = doc.part.blob.decode("utf-8")
    found_links = re.findall(r'href="(https?://[^"]+)"', xml_str)
    return list(set(found_links))

###############################################################################
# 13. WEIGHTED GROUPING FUNCTION FOR PDF/IMAGES
###############################################################################
def group_sections_by_header(sections, tables):
    """
    Groups text sections and tables by spatial proximity.
    For PDFs/images, each non-header section is assigned to the header above it.
    
    • For each non-header section, candidate headers above it are gathered.
    • If any candidate has an x difference within ±5% of the section's x, choose the candidate with the minimum weighted distance:
         distance = 0.3 * |sec.x - candidate.x| + 0.7 * (sec.y - candidate.y)
    • Otherwise, choose the candidate (ignoring x difference) with the minimum weighted distance.
    • If the vertical gap (sec.y - candidate.y) exceeds 50% of sec.y, assign the section to "PROFILE".
    """
    if isinstance(sections, dict):
        sections_list = list(sections.values())
    else:
        sections_list = sections
    sorted_sections = sorted(sections_list, key=lambda s: s.get("y", 0))
    groups = {}
    for sec in sorted_sections:
        if sec.get("is_header"):
            header_key = f"HEADER::{sec.get('canonical_header','')}::{sec.get('text','')}"
            groups.setdefault(header_key, {"header": sec, "text": [], "tables": []})
    for sec in sorted_sections:
        if sec.get("is_header"):
            continue
        x_val = sec.get("x", 0)
        y_val = sec.get("y", 0)
        candidates = [h for h in sorted_sections if h.get("is_header") and h.get("y", 0) < y_val]
        if candidates:
            x_tol_candidates = [c for c in candidates if abs(c.get("x", 0) - x_val) <= 0.05 * x_val]
            if x_tol_candidates:
                candidate = min(x_tol_candidates, key=lambda h: 0.3 * abs(h.get("x", 0)-x_val) + 0.7 * (y_val - h.get("y", 0)))
            else:
                candidate = min(candidates, key=lambda h: 0.3 * abs(h.get("x", 0)-x_val) + 0.7 * (y_val - h.get("y", 0)))
            vertical_gap = y_val - candidate.get("y", 0)
            if vertical_gap > 0.5 * y_val:
                sec["group"] = "PROFILE"
                groups.setdefault("PROFILE", {"header": None, "text": [], "tables": []})
                groups["PROFILE"]["text"].append(sec)
            else:
                header_key = f"HEADER::{candidate.get('canonical_header','')}::{candidate.get('text','')}"
                sec["group"] = header_key
                groups.setdefault(header_key, {"header": candidate, "text": [], "tables": []})
                groups[header_key]["text"].append(sec)
        else:
            sec["group"] = "PROFILE"
            groups.setdefault("PROFILE", {"header": None, "text": [], "tables": []})
            groups["PROFILE"]["text"].append(sec)
    tables = assign_table_header(tables)
    for tbl in tables:
        if tbl.get("detected_header"):
            header_key = f"HEADER::{tbl['detected_header']}::(table header)"
            matched = False
            for key in groups.keys():
                if tbl["detected_header"].lower() in key.lower():
                    tbl["group"] = key
                    groups[key]["tables"].append(tbl)
                    matched = True
                    break
            if not matched:
                tbl["group"] = header_key
                groups.setdefault(header_key, {"header": None, "text": [], "tables": []})
                groups[header_key]["tables"].append(tbl)
        else:
            bbox = tbl.get("bbox")
            if bbox:
                table_x, table_y, w, h = bbox
                candidates = [h for h in sorted_sections if h.get("is_header") and h.get("y", 0) < table_y]
                if candidates:
                    x_tol_candidates = [c for c in candidates if abs(c.get("x", 0) - table_x) <= 0.05 * table_x]
                    if x_tol_candidates:
                        candidate = min(x_tol_candidates, key=lambda h: 0.3 * abs(h.get("x", 0)-table_x) + 0.7 * (table_y - h.get("y", 0)))
                    else:
                        candidate = min(candidates, key=lambda h: 0.3 * abs(h.get("x", 0)-table_x) + 0.7 * (table_y - h.get("y", 0)))
                    vertical_gap = table_y - candidate.get("y", 0)
                    if vertical_gap > 0.5 * table_y:
                        tbl["group"] = "PROFILE"
                        groups.setdefault("PROFILE", {"header": None, "text": [], "tables": []})
                        groups["PROFILE"]["tables"].append(tbl)
                    else:
                        header_key = f"HEADER::{candidate.get('canonical_header','')}::{candidate.get('text','')}"
                        tbl["group"] = header_key
                        groups.setdefault(header_key, {"header": candidate, "text": [], "tables": []})
                        groups[header_key]["tables"].append(tbl)
                else:
                    tbl["group"] = "PROFILE"
                    groups.setdefault("PROFILE", {"header": None, "text": [], "tables": []})
                    groups["PROFILE"]["tables"].append(tbl)
            else:
                tbl["group"] = "PROFILE"
                groups.setdefault("PROFILE", {"header": None, "text": [], "tables": []})
                groups["PROFILE"]["tables"].append(tbl)
    return groups

###############################################################################
# 7D. ASSIGN TABLE HEADER FUNCTION
###############################################################################
def assign_table_header(tables):
    """
    For each table, attempts to detect a header using table cells.
    If detected, assigns it under "detected_header".
    """
    for tbl in tables:
        detected = detect_headers_in_tables([tbl])
        if detected:
            tbl["detected_header"] = detected[0]
        else:
            tbl["detected_header"] = None
    return tables

def detect_headers_in_tables(tables):
    """
    For each table, examines the header row (first row) and the left column.
    Returns a list of header texts detected in the tables.
    """
    table_headers = []
    for table in tables:
        data = table.get("data", [])
        if data:
            header_row = data[0]
            for cell in header_row:
                if cell and is_header_candidate(cell)[0]:
                    table_headers.append(cell.strip())
                    break
            for row in data:
                if row and row[0] and is_header_candidate(row[0])[0]:
                    table_headers.append(row[0].strip())
                    break
    return table_headers

###############################################################################
# 12. PROFILE EXTRACTION FROM THE PROFILE SECTION (Basic)
###############################################################################
def extract_profile(text):
    """
    Extracts profile fields (name, phone, email, GitHub, LinkedIn) from text using spaCy NER and regex.
    Returns a dictionary with the first detected name and other fields.
    The remaining text (with these fields removed) is stored under "data".
    """
    import re
    doc = nlp(text)
    
    email_match = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    phone_match = re.search(r'\+?\d[\d\-\s]{8,}\d', text)
    github_match = re.search(r'https?://(?:www\.)?github\.com/\S+', text)
    linkedin_match = re.search(r'https?://(?:www\.)?linkedin\.com/in/\S+', text)
    
    email = email_match.group(0) if email_match else ""
    phone = phone_match.group(0) if phone_match else ""
    github = github_match.group(0) if github_match else ""
    linkedin = linkedin_match.group(0) if linkedin_match else ""
    
    name = ""
    for ent in doc.ents:
        if ent.label_ == "PERSON":
            name = ent.text
            break
    
    data_text = text
    for field in [email, phone, github, linkedin, name]:
        if field:
            data_text = re.sub(re.escape(field), "", data_text, flags=re.IGNORECASE)
    data_text = re.sub(r'\s+', ' ', data_text).strip()
    
    return {
        "name": name or "Your Name",
        "phone": phone or "Phone",
        "mail": email or "yourmail@gmail.com",
        "github": github or "",
        "linkedin": linkedin or "",
        "data": data_text
    }

###############################################################################
# 12A. PROFILE DETAILS EXTRACTION USING NER AND REGEX
###############################################################################
def extract_profile_details(profile_text):
    """
    Extracts profile details from the given text.
    Uses spaCy NER to extract all PERSON entities (names) and regex to extract emails, phone numbers, and links.
    Returns a dictionary with lists of detected values.
    """
    import re
    doc = nlp(profile_text)
    
    names = list({ent.text for ent in doc.ents if ent.label_ == "PERSON"})
    emails = re.findall(r'[\w\.-]+@[\w\.-]+\.\w+', profile_text)
    phone_numbers = re.findall(r'\+?\d[\d\-\s]{8,}\d', profile_text)
    links = re.findall(r'https?://\S+', profile_text)
    
    return {
        "names": list(set(names)),
        "emails": list(set(emails)),
        "phone_numbers": list(set(phone_numbers)),
        "links": list(set(links))
    }

###############################################################################
# 12B. EMBEDDED LINK EXTRACTION
###############################################################################
def extract_embedded_links_from_pdf(pdf_path):
    """
    Extracts embedded hyperlinks from a PDF using PyMuPDF.
    Returns a list of unique URLs found in the link annotations.
    """
    import fitz
    doc = fitz.open(pdf_path)
    links = []
    for page in doc:
        for lnk in page.get_links():
            if "uri" in lnk:
                links.append(lnk["uri"])
    return list(set(links))

def extract_embedded_links_from_docx(docx_path):
    """
    Extracts embedded hyperlinks from a DOCX file.
    Retrieves the underlying XML and uses regex to find href attributes.
    """
    doc = Document(docx_path)
    xml_str = doc.part.blob.decode("utf-8")
    found_links = re.findall(r'href="(https?://[^"]+)"', xml_str)
    return list(set(found_links))

###############################################################################
# 13. RENDERING GROUPED SECTIONS
###############################################################################
def render_grouped_sections(grouped):
    """Prints each header group and its text (and table info) in a readable format."""
    print("\n==========================")
    print(" Final Grouped Sections ")
    print("==========================")
    for header_key, group in grouped.items():
        header = group.get("header")
        header_text = header.get("text", "") if header is not None else "PROFILE (No Header)"
        print(f"\n--- Header: {header_text} ---")
        for sec in group.get("text", []):
            print(f"  > {sec.get('text', '').strip()}")
        if group.get("tables"):
            print("  [Tables:]")
            for tbl in group.get("tables", []):
                page = tbl.get("page", "N/A")
                bbox = tbl.get("bbox", "")
                print(f"    - Table (Page: {page}): {bbox}")
    print("\n==========================\n")

###############################################################################
# 14. MAIN RUNNER FUNCTION
###############################################################################
def run_parser():
    """
    Overall Process:
      A. For PDFs/Images:
         - Convert the document to an image (first page for PDFs/images).
         - Extract text regions using a heatmap approach (starting with dilation radius 30).
         - Run OCR and header detection (including fallback for first 4 words).
         - Apply DBSCAN clustering to merge nearby fragmented bounding boxes.
         - Group sections by spatial proximity using a weighted distance metric (vertical threshold 50%).
      B. For DOCX files:
         - Extract paragraphs and group them sequentially.
      C. Render all grouped sections.
      D. Any section (or default group) that belongs to profile (i.e. the group key is "PROFILE" or its header text contains "profile")
         is processed using our NER and regex extraction to get name, email, phone, links, and summary.
         Also, if the file is a PDF or DOCX, embedded links are extracted.
         The final output is shown in JSON format.
    """
    file_path = input("\n📂 Enter File Path (PDF/DOCX/Image): ").strip()
    if not os.path.exists(file_path):
        print("❌ Error: File not found!")
        return

    is_docx = file_path.endswith(".docx")
    is_pdf = file_path.endswith(".pdf")
    
    if is_docx:
        sections = extract_text_from_docx(file_path)
        tables = extract_tables_from_docx(file_path)
        grouped = group_docx_sections(sections, tables)
    else:
        sections = process_file(file_path)
        tables = []
        if is_pdf:
            tables = extract_tables_from_pdf(file_path)
        elif file_path.lower().endswith((".png", ".jpg", ".jpeg", ".tiff", ".bmp")):
            tables = extract_tables_from_image(file_path)
        if tables:
            tables = assign_table_header(tables)
            detected_table_headers = detect_headers_in_tables(tables)
            if detected_table_headers:
                print(f"✅ Detected table headers: {detected_table_headers}")
        else:
            print("\nℹ️ No tables detected or found.")
        grouped = group_sections_by_header(list(sections.values()) if isinstance(sections, dict) else sections, tables)
    
    hdr_count = count_headers(sections)
    print(f"\n✅ Final header count from text regions: {hdr_count}")
    render_grouped_sections(grouped)
    
    # Find the group(s) that belong to Profile.
    # We consider a group as Profile if its key is "PROFILE" or if its header text contains "profile".
    profile_group = None
    for key, group in grouped.items():
        if key.lower() == "profile" or ("profile" in key.lower()):
            profile_group = group
            break

    if profile_group is not None:
        profile_text = "\n".join(sec.get("text", "") for sec in profile_group.get("text", []))
        basic_profile = extract_profile(profile_text)
        detailed_profile = extract_profile_details(profile_text)
        embedded_links = []
        if is_pdf:
            embedded_links = extract_embedded_links_from_pdf(file_path)
        elif is_docx:
            embedded_links = extract_embedded_links_from_docx(file_path)
        
        # Merge regex-extracted links with embedded links.
        all_links = list(set(detailed_profile.get("links", []) + embedded_links))
        detailed_profile["embedded_links"] = all_links
        
        final_profile = {
            "name": basic_profile.get("name"),
            "email": basic_profile.get("mail"),
            "phone": basic_profile.get("phone"),
            "github": basic_profile.get("github"),
            "linkedin": basic_profile.get("linkedin"),
            "summary": basic_profile.get("data"),
            "regex_links": detailed_profile.get("links"),
            "embedded_links": detailed_profile.get("embedded_links")
        }
        
        print("\n=== Extracted Profile Data ===")
        print(json.dumps(final_profile, indent=2))
    else:
        print("\n⚠️ No Profile section found.")

# === RUN THE SCRIPT ===
if __name__ == "__main__":
    run_parser()
